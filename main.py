import os
import tkinter as tk
from tkinter import filedialog, font
from docx import Document
from datetime import datetime

class HereApp:
    def __init__(self, window):
        self.window = window
        self.window.title('Here App')
        self.window.geometry('520x480')

        label_font = font.Font(family='Helvetica', size=14, weight='bold')
        self.label = tk.Label(window, text='Upload Class Attendance Sheet', height=3, width=40, font=label_font)
        self.label.pack(expand=True, padx=15, pady=10)

        student_font = font.Font(family='Helvetica', size=28, weight='bold')
        self.student_label = tk.Label(window, text='', font=student_font, wraplength=400)
        self.student_label.pack(padx=5, pady=5)

        button_font = font.Font(family='Helvetica', size=14, weight='bold')
        self.present_button = tk.Button(window, text='Present', command=lambda: self.mark_attendance(True), height=3, width=10, bg='green', font=button_font)
        self.present_button.pack(side=tk.LEFT, expand=True, padx=10, pady=50)
        self.absent_button = tk.Button(window, text='Absent', command=lambda: self.mark_attendance(False), height=3, width=10, bg='red', font=button_font)
        self.absent_button.pack(side=tk.RIGHT, expand=True, padx=10, pady=50)

        self.upload_button = tk.Button(window, text='Upload Document', command=self.upload_document, height=3, width=15, font=button_font)
        self.upload_button.pack(side=tk.BOTTOM, pady=5)

        self.name_label = tk.Label(window, text='made by jude :)')
        self.name_label.place(relx=1.0, rely=1.0, x=-10, y=-10, anchor='se')

        self.doc = None
        self.table = None
        self.doc_path = None
        self.student_row_number = 0

    def run(self):
        self.update_label('Upload Attendance Document')
        self.window.mainloop()

    def get_formatted_today_date(self):
        # Get today's date in the format "Jan 1" (with leading zero removed)
        return datetime.now().strftime("%b %d").replace(" 0", " ")

    def find_today_date_column(self):
        today_date_str = self.get_formatted_today_date()
        uppercased_today_date_str = today_date_str.upper()

        header_row = self.table.rows[0]
        for date, cell in enumerate(header_row.cells):
            cell_text = cell.text.strip()

            if today_date_str in cell_text or uppercased_today_date_str in cell_text:
                return date
        return None

    def action_on_upload(self, file_path):
        if not (file_path.endswith('.doc') or file_path.endswith('.docx')):
            self.update_label('Wrong Document type uploaded')
            return

        try:
            self.doc = Document(file_path)
            self.doc_path = file_path

            if not self.doc.tables:
                self.update_label('No tables found in document')
                return

            self.table = self.doc.tables[0]
        except Exception as e:
            self.update_label(f'Error loading document: {str(e)}')
            return

        self.date_column_number = self.find_today_date_column()

        if self.date_column_number is not None:
            self.student_row_number = 1
            self.update_student_label()
        else:
            formatted_date = self.get_formatted_today_date()
            self.update_label(f"Date {formatted_date} not found in document")

    def action_on_button_press(self, is_present):
        if self.student_row_number < len(self.table.rows):
            if not is_present:
                self.table.cell(self.student_row_number, self.date_column_number).text = 'Absent'

            self.student_row_number += 1
            self.update_student_label()

    def update_student_label(self):
        if self.student_row_number < len(self.table.rows):
            student_name = self.table.cell(self.student_row_number, 0).text
            self.update_label(student_name)
        else:
            self.save_document()
            self.update_label('Attendance Complete')

    def save_document(self):
        if self.doc_path:
            self.doc.save(self.doc_path)
            self.update_label('Document saved')

    def set_document_name(self, file_path):
        document_name = os.path.basename(file_path)
        document_name_no_extension = os.path.splitext(document_name)[0]
        self.label.config(text=document_name_no_extension)

    def upload_document(self):
        file_path = filedialog.askopenfilename(filetypes=[('Word Documents', '*.docx'), ('Word Document', '.doc')])
        if file_path:
            self.action_on_upload(file_path)
            self.set_document_name(file_path)

    def update_label(self, text):
        self.student_label.config(text=text)

    def mark_attendance(self, is_present):
        self.action_on_button_press(is_present)


if __name__ == '__main__':
    window = tk.Tk()
    here_app = HereApp(window)
    here_app.run()
